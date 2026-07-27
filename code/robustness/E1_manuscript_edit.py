"""E1 + reframe compile: apply the D1 placebo-test outcome to FINAL_PAPER_V25_PLOS.docx.

Withdraws the Gibbs-specific/structural interpretation of the asymmetric
temperature prediction (§4.4's "strongest and most credible evidence" claim).
Retains, restated, the residual full-universe conditional-pricing finding that
survives dispersion normalization. Adds the placebo-characteristic table.
Reframes title/abstract/conclusion around the survivorship correction as the
paper's actual contribution, per the Decision Point.

Base: FINAL_PAPER_V25_PLOS.docx (no V26/V27 draft exists in this repo).
Output: FINAL_PAPER_V26_PLOS.docx
"""
import docx
import copy

SRC = "FINAL_PAPER_V25_PLOS.docx"
OUT = "FINAL_PAPER_V26_PLOS.docx"

doc = docx.Document(SRC)
paras = doc.paragraphs

def set_text(p, new_text):
    """Replace a paragraph's text, collapsing to a single run."""
    if not p.runs:
        p.add_run(new_text)
        return
    p.runs[0].text = new_text
    for r in list(p.runs[1:]):
        r._element.getparent().remove(r._element)

def append_text(p, addition):
    if not p.runs:
        p.add_run(addition)
        return
    p.runs[-1].text = p.runs[-1].text.rstrip() + " " + addition

# ═════════════════════════════════════════════════════════════════════════
# PHASE 1 — in-place text replacements (indices stable, no insertions yet)
# ═════════════════════════════════════════════════════════════════════════

# -- Title --
set_text(paras[0],
    "Survivorship Bias and the Idiosyncratic Volatility Premium: "
    "A Gibbs Free Energy-Motivated Decomposition of Quality and Disorder")

# -- Abstract --
set_text(paras[3],
    "We apply the Gibbs free energy equation (ΔG = ΔH − TΔS) as a motivating "
    "decomposition for cross-sectional equity return prediction, mapping gross profit margin "
    "stability to enthalpy (ΔH), idiosyncratic volatility to entropy (ΔS), and realized "
    "market variance to temperature (T). The paper’s central finding is that the apparent "
    "positive relationship between a long-horizon idiosyncratic-volatility measure and returns "
    "in large-cap survivor panels is a survivorship artifact: in a survivorship-corrected "
    "full-universe SF1 panel of 12,449 firms (72% subsequently delisted), the disorder premium "
    "collapses from FM t = +4.70 to +0.02 at quarterly frequency. Two tests support the "
    "attribution: rebuilding ΔS under the identical quarterly construction on the "
    "survivorship-biased S&P 500 panel leaves the premium intact (FM t = +3.65), and restricting "
    "the corrected panel to the largest 500 firms with delisted names retained still finds none "
    "(FM t = +0.58) — eliminating measurement and size composition as explanations. A "
    "synthetic stress test eliminates the +13.4%/yr quintile premium at delisting rates implied "
    "by observed performance-delisting frequencies; survival conditioning resurrects it "
    "monotonically, reaching FM t = +3.23 among 27-year survivors. Prior positive entropy-return "
    "claims, including Fu (2009) via look-ahead bias and Ormos and Zibriczky (2014) via 27-year "
    "survival conditioning, share this sample-construction class. Second, a quality premium "
    "survives: the gross profit margin stability channel strengthens to FM t = +3.46, clearing "
    "the Harvey-Liu-Zhu threshold in the sample that includes failed firms. A secondary "
    "conditional pattern — the entropy loading co-varies positively with market temperature "
    "(HAC t = +2.70, survivorship-corrected) while the stability loading does not — "
    "motivated the Gibbs equation asymmetric prediction, but a placebo test regressing size, "
    "book-to-market, momentum, and beta through the identical design finds the same "
    "temperature-covariation in most of them; we therefore treat the pattern as generic "
    "conditional pricing rather than evidence for the thermodynamic mapping.")

# -- Intro paragraph [13] --
set_text(paras[13],
    "A secondary conditional pattern, reported here as directional rather than established, "
    "appears in both samples: the entropy loading co-varies positively with market temperature "
    "(HAC t = +2.70, survivorship-corrected full-universe panel) while the stability loading "
    "does not (HAC t = −1.51). We initially read this asymmetry as evidence for the Gibbs "
    "equation’s structural form, since ΔH enters the equation unscaled and ΔS "
    "scaled by T. A placebo test (Section 4.4) regressing four characteristics unrelated to "
    "disorder — size, book-to-market, momentum, and market beta — through the "
    "identical first-step-then-HAC design finds comparable temperature-covariation in three of "
    "four placebo characteristics in both panels, and in all four in at least one panel, because "
    "cross-sectional return dispersion itself rises with T and inflates any characteristic’s "
    "second-step loading. The asymmetry is therefore not distinctive to entropy and does not by "
    "itself support the thermodynamic mapping specifically; we retain it as a generic "
    "conditional-pricing observation. It also weakens in the post-2009 subsample (cluster-robust "
    "Wald p ≈ 0.07) and is statistically on the boundary in bootstrap (median p = 0.063, "
    "42.5% of samples below p = 0.05). One identification issue applies to the FM estimator: in "
    "the encompassing model, T·ΔS_i,t is collinear with ΔS_i,t within each "
    "cross-section because T is constant across stocks, so the encompassing FM cannot separately "
    "estimate a T·ΔS coefficient. Model C (ΔH + T·ΔS, substituting "
    "T·ΔS for ΔS) does not have this collinearity and yields FM t = +4.59, though "
    "as Section 4.3 explains this statistic reflects level pricing of ΔS rather than "
    "T-scaling per se; the estimators that identify T-scaling are the cluster-robust Wald "
    "(p = 0.013), the pooled two-way-clustered interaction (t = +2.49 in both panels), and the "
    "asymmetric HAC test (Section 4.4), none of which by themselves establish that the T-scaling "
    "is specific to disorder rather than a generic feature of conditional dispersion.")

# -- §2.4 Temperature Amplification: append clarifying sentence --
append_text(paras[34],
    "A related concern, addressed directly in Section 4.4’s placebo battery, is that any "
    "characteristic’s second-step loading can appear to co-vary with T simply because "
    "cross-sectional return dispersion rises with market volatility, independent of whether "
    "disorder specifically is rewarded; that placebo test — not H3 — is the more "
    "direct test of whether the entropy-temperature link is thermodynamically distinctive rather "
    "than a mechanical consequence of dispersion scaling.")

# -- §2.5 H2 statement --
set_text(paras[38],
    "H2 (T-Scaling): The temperature interaction T·ΔS is significant in "
    "specifications that identify it — the cluster-robust Wald test, the pooled "
    "two-way-clustered interaction, and the asymmetric HAC test — and, if the Gibbs mapping "
    "is thermodynamically specific rather than a generic dispersion-scaling channel, this "
    "T-covariation should be distinctive to the entropy channel and largely absent from "
    "unrelated characteristics. Outcomes are reported in Sections 4.3, 4.4, and 4.7: the plain "
    "significance claim holds (Wald p = 0.013/0.017; pooled interaction t = +2.49 both panels; "
    "asymmetric HAC t = +2.25 in-sample / +2.70 survivorship-corrected). The distinctiveness "
    "sub-claim does not hold: a placebo test of size, book-to-market, momentum, and beta through "
    "the identical design (Section 4.4) finds comparable temperature-covariation in most of "
    "them, consistent with a generic dispersion-scaling mechanism rather than a channel specific "
    "to disorder. (The encompassing-model FM is uninformative for T·ΔS due to "
    "within-cross-section collinearity, and Model C FM does not by itself isolate T-scaling from "
    "level pricing of ΔS; see Section 4.3.)")

# -- §2.5 Falsifiability criterion: add third bullet --
set_text(paras[40],
    "Graded Falsifiability Criterion: "
    "- Strong falsification: The cluster-robust Wald test for T·ΔS fails to achieve "
    "p < 0.05 in the full-sample pooled specification under both constructions. This would imply "
    "the thermodynamic organizing structure has no empirical support. "
    "- Partial falsification: The T-scaling result fails in more than half of subperiod "
    "exclusion windows under both constructions simultaneously. The post-2009 weakness "
    "(p ≈ 0.07, both constructions) constitutes a partial qualification: the framework "
    "applies more reliably to high-temperature-variance regimes than to the "
    "compressed-volatility post-2010 environment. "
    "- Distinctiveness falsification (added post hoc, following referee review): if two or more "
    "of four placebo characteristics unrelated to disorder (size, book-to-market, momentum, "
    "beta) show second-step temperature-covariation of comparable magnitude and significance to "
    "ΔS under the identical design, the asymmetric T-scaling result cannot be attributed to "
    "the Gibbs mapping specifically and must be treated as a generic conditional-pricing / "
    "dispersion-scaling pattern rather than structural confirmation of the thermodynamic form.")

set_text(paras[41],
    "Strong falsification does not apply: the Wald test clears p < 0.05 in both constructions. "
    "Partial falsification applies in a limited sense via the post-2009 weakness. "
    "Distinctiveness falsification applies: four of four placebo characteristics show "
    "significant temperature-covariation in at least one panel (three of four in both), so the "
    "asymmetric T-scaling result is downgraded from a structural, thermodynamically-specific "
    "claim to a directional conditional-pricing observation consistent with, but not diagnostic "
    "of, the Gibbs functional form. This is the paper’s most consequential post hoc "
    "falsification result and the primary reason for reframing the paper’s contribution "
    "around the survivorship correction rather than the thermodynamic mapping (Section 5.1).")

# -- §4.4 paragraph [93]: soften "primary structural claim" framing --
set_text(paras[93],
    "The cluster-robust Wald test establishes that T·ΔS is significant, but a "
    "significant temperature interaction alone does not distinguish the Gibbs framework from a "
    "generic regime-conditional pricing model. Any model in which the iVol premium varies with "
    "market volatility would predict a significant T·ΔS term. The Gibbs equation makes "
    "a sharper, asymmetric prediction that a generic model does not: because enthalpy and "
    "entropy enter the free-energy equation differently, with ΔH unscaled and ΔS "
    "multiplied by T, the financial mapping implies that the enthalpy loading should be "
    "temperature-independent while the entropy loading should scale with temperature. "
    "Establishing this asymmetry, however, is not sufficient by itself to attribute it to the "
    "Gibbs mapping specifically, since cross-sectional return dispersion also rises with market "
    "volatility and could inflate any characteristic’s estimated second-step loading "
    "regardless of its economic content. We test the asymmetry directly below, and then test its "
    "distinctiveness with a placebo battery of characteristics unrelated to disorder.")

# -- §4.4 paragraph [96]: remove the withdrawn "strongest and most credible evidence" claim --
set_text(paras[96],
    "Both halves of the asymmetric prediction are consistent with the data in both samples: T "
    "enters through the entropy channel, and no temperature dependence is detected in the "
    "enthalpy channel. One caution applies to the second half: the enthalpy result is a failure "
    "to reject rather than a demonstrated null, and the difference between a significant and an "
    "insignificant slope is not itself a test of their difference (Gelman and Stern, 2006); we "
    "therefore implement a direct test of the slope difference here, stacking the two second-step "
    "series and interacting temperature with a channel indicator, with standard errors clustered "
    "by period. A direct test of the slope difference confirms the asymmetry rather than leaving "
    "it as a contrast between a significant and an insignificant slope: interacting temperature "
    "with a channel indicator on the stacked second-step series yields a positive, significant "
    "temperature-by-entropy interaction in both panels (full universe: +0.60, date-clustered "
    "t = +2.26, p = 0.024; S&P 500: +0.15, t = +2.23, p = 0.026), and the entropy channel’s "
    "temperature-dependence strengthens under survivorship correction (HAC t = +2.70 versus "
    "+2.25 in-sample). This establishes that the asymmetry is real and not an artifact of "
    "comparing a significant to an insignificant slope. It does not, on its own, establish that "
    "the asymmetry reflects the Gibbs mapping specifically rather than a mechanical consequence "
    "of dispersion scaling with T — a distinct question we address directly next.")

# -- §4.7: point to the placebo test --
append_text(paras[106],
    "Section 4.4 reports a placebo-characteristic test of this asymmetry; three or four of four "
    "characteristics unrelated to disorder show comparable temperature-covariation, so the "
    "significance results above should not be read as diagnostic of the Gibbs mapping "
    "specifically.")

set_text(paras[107],
    "Bootstrap distribution of T·ΔS Wald. Over 1,000 block-bootstrap samples (block "
    "length = 12 months): mean t = +1.83, median p = 0.063, 42.5% of samples achieve p < 0.05, "
    "66.4% achieve p < 0.10, and 99.5% are positive (t > 0). The 5th/50th/95th percentiles of the "
    "t-distribution are +0.84/+1.86/+2.70. Sign certainty is essentially complete (99.5%); the "
    "direction of H3 — |β_ΔS| larger in the high-T than the low-T regime, from "
    "separate regime FM regressions within each bootstrap sample — holds in 83.4% of "
    "samples. Statistical significance at p < 0.05 is marginal. The most accurate summary, "
    "updated for the placebo test in Section 4.4: directionally certain in sign, statistically "
    "on the boundary, and not distinctive to the entropy channel relative to size, "
    "book-to-market, and momentum — a generic conditional-pricing pattern rather than "
    "confirmed temperature amplification specific to disorder.")

# -- §5.1 --
set_text(paras[148],
    "The paper’s arc is the progressive removal of artifacts, including from its own "
    "central organizing device. The price-based ΔG composite was a disguised negative iVol "
    "factor (Corr = −0.853); the accounting rebuild corrected it. The in-sample disorder "
    "premium (FM t = +4.70) is largely a survivorship artifact; the R18/R19 corrections "
    "established that, and the survival-conditioning ladder (R20, Table 8) shows the artifact "
    "can be manufactured to order by the sample design alone. A placebo-characteristic test "
    "(Section 4.4) removes a further piece: the asymmetric temperature prediction — "
    "ΔS loading on T, ΔH not — is not distinctive to disorder, since size, "
    "book-to-market, and momentum show comparable temperature-covariation under the identical "
    "design. What remains after all three corrections is: a confirmed quality premium in the "
    "stability channel (β_ΔH, FM t = +3.46, HLZ-clearing, full-universe), and, from the "
    "entropy channel, only a directional and non-thermodynamically-specific conditional pattern "
    "(β_ΔS ~ T, HAC t = +2.70 raw / +2.57 dispersion-normalized in the full-universe "
    "panel, fragile by bootstrap, weaker post-2009, and generic across unrelated characteristics "
    "rather than diagnostic of the Gibbs mapping).")

set_text(paras[151],
    "What the survivorship correction does not eliminate, and what the placebo test does not "
    "fully explain away either, is the T-scaling pattern in the full-universe panel: "
    "β_ΔS,t co-varies positively with T_t even when the average β_ΔS ≈ 0, "
    "and part of that covariation survives dispersion normalization (HAC t = +2.57). This is "
    "coherent — the entropy loading is positive in high-T months and near-zero or negative "
    "in low-T months, canceling on average but correlating with temperature — but we no "
    "longer read it as confirmation of the Gibbs equation’s specific functional form, since "
    "the placebo test shows the same directional pattern in characteristics the Gibbs mapping "
    "says nothing about. The most defensible statement is narrower than the one earlier drafts "
    "of this paper made: idiosyncratic-volatility pricing in this sample has a conditional, "
    "temperature-dependent component that is not purely a mechanical consequence of rising "
    "cross-sectional dispersion, but neither its magnitude nor its channel-specificity supports "
    "treating it as evidence for a thermodynamic organizing structure. We report it as a "
    "directional finding motivating future work on the mechanism, rather than as a confirmed "
    "structural feature of the Gibbs mapping.")

# -- §6 Conclusion --
set_text(paras[179],
    "A secondary conditional pattern is reported as directional rather than established. The "
    "entropy loading co-varies positively with market temperature in both samples (HAC t = "
    "+2.70, survivorship-corrected full-universe; +2.25 in-sample) while the stability loading "
    "does not (HAC t = −1.51). We initially read this asymmetry as consistent with the "
    "functional form ΔG = ΔH − TΔS. A placebo-characteristic test reported "
    "in Section 4.4 shows the same temperature-covariation in most of four unrelated "
    "characteristics (size, book-to-market, momentum, and, in one panel, beta), so the asymmetry "
    "is not distinctive to disorder and does not support the Gibbs mapping specifically; part of "
    "it survives dispersion-normalization in the full-universe panel (HAC t = +2.57), so we "
    "retain it as a generic, directional conditional-pricing observation rather than a confirmed "
    "thermodynamic structural feature.")

set_text(paras[180],
    "The paper’s contribution is accurately bounded, and narrower than earlier drafts of "
    "this paper claimed. It does not recover a positive disorder premium that survives bias "
    "correction, and it does not establish that idiosyncratic-volatility pricing is "
    "thermodynamically structured: the encompassing ΔG composite is insignificant, the "
    "temperature-amplification hypothesis (H3) is unsupported, the literal unit-ratio constraint "
    "is untestable given the sample’s power, and the asymmetric temperature prediction "
    "— the paper’s last Gibbs-specific empirical claim — does not survive a "
    "placebo test showing the same pattern in unrelated characteristics. What the paper does "
    "establish is a survivorship-correction result and a quality-premium result, both robust: it "
    "documents the mechanism by which iVol premiums are inflated in large-cap samples, "
    "quantifies the inflation, and confirms that earnings stability is robustly priced across "
    "sample constructions, including one that retains the firms that failed. The Gibbs free "
    "energy equation motivated the decomposition into a stability channel and a disorder channel "
    "that made this quantification possible; that motivating role, rather than a confirmed "
    "physical analogy, is the paper’s honest relationship to the thermodynamic framework.")

print("Phase 1 (in-place replacements) complete.")

# ═════════════════════════════════════════════════════════════════════════
# PHASE 2 — insertion: placebo-battery prose + table, after paragraph [96]
# ═════════════════════════════════════════════════════════════════════════
p97 = doc.paragraphs[97]  # measurement-error note; insert before this

textA = (
    "Placebo-characteristic test. β̂_ΔS,t is a cross-sectional slope on a "
    "standardized regressor, proportional to cov(r_{i,t+1}, ΔS^z_{i,t}); since "
    "cross-sectional return dispersion rises with market volatility (Corr(σ_cs,t+1, T_t) = "
    "+0.39 in-sample, +0.47 survivorship-corrected), any characteristic with a non-zero mean "
    "loading could produce a second-step slope that co-varies with T mechanically, with no "
    "conditional-pricing content whatsoever. To test whether the entropy-temperature link is "
    "distinctive, we re-run the identical first-step-then-HAC design (same panel, same controls, "
    "same months, β̂_char,t regressed on T_t) for four characteristics unrelated to "
    "disorder: size (log market capitalization), book-to-market, momentum (12-1), and market "
    "beta. Table 4.4a reports the results: three of four placebo characteristics (size, "
    "book-to-market, momentum) show HAC t > 2 in both panels; the fourth (beta) is significant "
    "in the S&P 500 panel (HAC t = +2.61) but not in the full-universe panel (HAC t = +1.25). "
    "Four of four placebo characteristics are significant on T in at least one panel."
)

caption = (
    "Table 4.4a: Placebo-Characteristic Second-Step Test, β̂_char,t ~ T_t "
    "(final table number pending the Table renumbering pass, Section E3)"
)

textB = (
    "Normalizing the entropy loading by the same-period cross-sectional standard deviation of "
    "the priced return (β̂_ΔS,t / σ_cs,t+1) partially separates the "
    "mechanical dispersion channel from a genuine conditional-pricing signal. In the S&P 500 "
    "in-sample panel, the normalized HAC t falls from +2.25 (raw) to +1.39, below conventional "
    "significance. In the survivorship-corrected full-universe panel — the panel this "
    "paper’s headline asymmetric-prediction claim is built on — the normalized HAC t "
    "is +2.57 (p = 0.010), close to the raw +2.70 and still significant. The stacked "
    "date-clustered slope-difference test (the direct test of asymmetry introduced above) shows "
    "the same pattern when re-run on normalized loadings: the full-universe difference remains "
    "significant (d = +1.65, t = +2.21, p = 0.027) while the S&P 500 difference does not "
    "(d = +0.74, t = +1.25, p = 0.213). We draw two conclusions. First, on the placebo count "
    "alone, the asymmetric temperature prediction cannot be treated as distinctive evidence for "
    "the thermodynamic mapping: loading on T is the generic case across the characteristics "
    "tested, not a property unique to entropy, so we withdraw the earlier claim that this "
    "asymmetry is the paper’s strongest and most credible evidence that the thermodynamic "
    "mapping captures something structural. Second, the dispersion-normalized result in the "
    "survivorship-corrected full-universe panel — the sample the paper’s headline "
    "claim is drawn from — remains significant after this correction, so there is a "
    "residual conditional relationship between the entropy loading and market temperature that "
    "is not purely mechanical. We report this as a directional conditional-pricing observation, "
    "not a thermodynamically specific one: something about the entropy channel’s pricing "
    "does vary with market temperature beyond what dispersion scaling alone predicts, but the "
    "placebo battery removes our basis for attributing that residual to the Gibbs functional "
    "form rather than to some other conditional-pricing mechanism operating on "
    "idiosyncratic-volatility-sorted portfolios generally. See Section 5.1 for the resulting "
    "reframing of the paper’s contribution."
)

pA = p97.insert_paragraph_before(textA)
pCap = p97.insert_paragraph_before(caption)
pB = p97.insert_paragraph_before(textB)

# style caption like other table captions if a "Table Caption"-ish style exists on a known table para
try:
    pCap.style = doc.styles['Caption']
except KeyError:
    pass

# Build the table and move it into place (right after the caption, before pB)
rows = [
    ("Characteristic", "S&P 500 HAC-t", "Full-Universe HAC-t", "Note"),
    ("Size (log mkt cap)", "−2.89", "−2.72", "significant, both panels"),
    ("Book-to-market", "+2.24", "+2.65", "significant, both panels"),
    ("Momentum (12-1)", "−2.15", "−4.00", "significant, both panels"),
    ("Market beta", "+2.61", "+1.25", "significant S&P only"),
    ("ΔS (paper's result, raw)", "+2.25", "+2.70", "reference row"),
    ("ΔS, dispersion-normalized", "+1.39", "+2.57", "S&P drops below significance; full-universe survives"),
    ("Stacked diff test, raw", "t=+2.23, p=0.026", "t=+2.26, p=0.024", "date-clustered"),
    ("Stacked diff test, normalized", "t=+1.25, p=0.213", "t=+2.21, p=0.027", "date-clustered"),
]
table = doc.add_table(rows=len(rows), cols=4)
try:
    table.style = doc.tables[0].style
except Exception:
    pass
for ri, row in enumerate(rows):
    for ci, val in enumerate(row):
        table.cell(ri, ci).text = val

pB._p.addprevious(table._tbl)

doc.save(OUT)
print(f"Saved: {OUT}")
print(f"Total paragraphs now: {len(doc.paragraphs)}  (was 220)")
print(f"Total tables now: {len(doc.tables)}  (was 18)")

abstract_wc = len(doc.paragraphs[3].text.split())
print(f"Abstract word count: {abstract_wc}")
