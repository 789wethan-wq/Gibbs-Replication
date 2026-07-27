"""build_SI.py — Block 4: Supporting Information exports (S1, S2) + M10-R note.

S1_Table  price-based composite ΔG portfolio sort (superseded, Section 4.1)
S2_Table  576-specification robustness battery (Section 4.7)
M10R_distinctness_note  per-panel pooled-Wald dimensions (closes the "two ≈2.49" question)
"""
import os
import numpy as np
import pandas as pd
import statsmodels.api as sm
from docx import Document

os.chdir(os.path.dirname(os.path.abspath(__file__)))
DATA = "../data"
DD = "../../Gibbs-Replication-sync/derived-data"
SI = "../../Gibbs-Replication-sync/supporting_information"
os.makedirs(SI, exist_ok=True)

FF = ["Mkt_RF", "SMB", "HML", "RMW", "CMA", "Mom"]


def nw_t(x, lags=6):
    x = np.asarray(x, float)
    x = x[~np.isnan(x)]
    n = len(x); mn = x.mean(); e = x - mn
    v = (e @ e) / n
    for L in range(1, lags + 1):
        v += 2 * (1 - L / (lags + 1)) * (e[L:] @ e[:-L]) / n
    return mn, mn / np.sqrt(v / n)


# ══ S1 — price-based ΔG sort ═══════════════════════════════════════════════════
pr = pd.read_csv(f"{DD}/portfolio_returns_monthly_sp500.csv", parse_dates=["date"])
fac = pd.read_parquet(f"{DATA}/factors_monthly.parquet")
fac.index = pd.to_datetime(fac.index)
fac = fac.reset_index().rename(columns={"index": "date"})
fac["date"] = pd.to_datetime(fac["date"])
d = pr.merge(fac, on="date", how="left")

rows = []
for q in ["DG_Q1", "DG_Q2", "DG_Q3", "DG_Q4", "DG_Q5", "DG_LS"]:
    ret = d[q].values
    mean_m, t_m = nw_t(ret * 100)
    y = (d[q] - d["RF"]) if q != "DG_LS" else d[q]  # L/S is already a spread
    reg = d[["date", q]].assign(y=y).dropna().merge(fac, on="date")
    X = sm.add_constant(reg[FF])
    res = sm.OLS(reg["y"], X).fit(cov_type="HAC", cov_kwds={"maxlags": 6})
    rows.append({
        "Quintile": q.replace("DG_", "").replace("LS", "Q5-Q1 (L/S)"),
        "Mean ret %/mo": round(mean_m, 3),
        "NW t (mean)": round(t_m, 2),
        "FF5+UMD alpha %/yr": round(res.params["const"] * 12 * 100, 2),
        "alpha NW t": round(res.tvalues["const"], 2),
    })
s1 = pd.DataFrame(rows)
s1.to_excel(f"{SI}/S1_Table.xlsx", index=False)

S1_CAP = ("S1 Table. Price-based composite portfolio sort (superseded construction, "
          "Section 4.1). Quintile-average monthly returns, the long-short spread, and "
          "Fama-French five-factor plus momentum (FF5+UMD) alphas for the price-based "
          "composite ΔG = ΔH_z - T·ΔS_z. Portfolios are formed monthly over the S&P 500 "
          "survivorship-biased panel, Jan 1995 - Nov 2023 (347 months). Mean-return "
          "t-statistics use Newey-West (6 lags); alphas are annualized with HAC(6) SEs. "
          "This construction is superseded by the accounting-based ΔH in the main text.")
doc = Document()
doc.add_paragraph(S1_CAP)
t = doc.add_table(rows=1, cols=len(s1.columns))
for j, c in enumerate(s1.columns):
    t.rows[0].cells[j].text = str(c)
for _, r in s1.iterrows():
    cells = t.add_row().cells
    for j, c in enumerate(s1.columns):
        cells[j].text = str(r[c])
doc.save(f"{SI}/S1_Table.docx")
print("S1 built:\n", s1.to_string(index=False))

# ══ S2 — 576-spec battery ══════════════════════════════════════════════════════
mrt = pd.read_csv(f"{DD}/master_robustness_table.csv")
S2_CAP = ("S2 Table. Full robustness battery (Section 4.7). Fama-MacBeth t-statistics, "
          "long-short returns, and model-comparison statistics across the "
          f"{len(mrt)}-specification grid. Columns: {', '.join(mrt.columns)}.")
with open(f"{SI}/S2_Table.csv", "w") as f:
    f.write("# " + S2_CAP + "\n")
    mrt.to_csv(f, index=False)
print(f"\nS2 built: {len(mrt)} specs -> S2_Table.csv (caption as leading # comment)")

# ══ M10-R — pooled-Wald distinctness note ══════════════════════════════════════
def cs_wz(df, col, dc="date", pct=0.01):
    def _f(x):
        x2 = x.dropna()
        if len(x2) < 5: return pd.Series(np.nan, index=x.index)
        lo, hi = x2.quantile(pct), x2.quantile(1 - pct)
        xc = x.clip(lo, hi); s = xc.std()
        return (xc - xc.mean()) / s if s > 1e-10 else pd.Series(np.nan, index=x.index)
    return df.groupby(dc)[col].transform(_f)

m = pd.read_parquet(f"{DATA}/merged_with_accounting.parquet"); m["date"] = pd.to_datetime(m["date"])
m["dH_gpm_z"] = cs_wz(m, "dH_gpm")
sp = m.dropna(subset=["ret_next_month", "dH_gpm_z", "DS_z", "TxDS"])
q = pd.read_parquet(f"{DATA}/merged_sf1_quarterly_survfree.parquet")
fu = q.dropna(subset=["ret_next", "delta_h_z", "delta_s_z", "T_delta_s"])

note = []
note.append("Supporting Note (M10-R). Distinctness of the two pooled T·ΔS Wald tests.")
note.append("")
note.append("The S&P 500 and full-universe pooled two-way-clustered Wald tests both land at "
            "t ≈ +2.49, but they are computed on different panels with different design "
            "matrices and cluster structures. They are not the same number reported twice.")
note.append("")
hdr = f"{'panel':<16}{'N obs':>10}{'T (date clus)':>15}{'firm clusters':>15}{'X shape':>12}{'coef T·ΔS':>11}"
note.append(hdr)
for lab, df, dcol, fcol, xc in [
    ("S&P 500", sp, "date", "stock_id", ["dH_gpm_z", "DS_z", "TxDS"]),
    ("Full-universe", fu, "q", "ticker", ["delta_h_z", "delta_s_z", "T_delta_s"]),
]:
    N = len(df); Tn = df[dcol].nunique(); Fn = df[fcol].nunique()
    Xshape = f"{N}x{len(xc)+1}"
    coef = float(np.linalg.lstsq(
        np.column_stack([np.ones(N)] + [df[c].values for c in xc]),
        df["ret_next_month" if dcol == "date" else "ret_next"].values, rcond=None)[0][-1])
    note.append(f"{lab:<16}{N:>10,}{Tn:>15}{Fn:>15,}{Xshape:>12}{coef:>+11.4f}")
note.append("")
note.append("Both pooled designs are full rank (4/4). See results/revision/PanelC_verify_ab.txt "
            "for the rank/condition/Wald detail. The shared t ≈ 2.49 is a coincidence of two "
            "independent estimates, not a duplicated cell.")
txt = "\n".join(note)
with open(f"{SI}/M10R_distinctness_note.txt", "w") as f:
    f.write(txt + "\n")
docn = Document()
for line in note:
    docn.add_paragraph(line)
docn.save(f"{SI}/M10R_distinctness_note.docx")
print("\n" + txt)
print(f"\nAll SI files written to {os.path.abspath(SI)}")
