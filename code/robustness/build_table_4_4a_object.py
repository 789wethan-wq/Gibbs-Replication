"""Build the actual Table 4.4a docx table object as a standalone, portable
file — for insertion into V27 (wherever it lives) since this session cannot
locate/access that file directly.

Contains: caption paragraph + the formatted Word table object with the D1
placebo-characteristic results. Matches the exact data already handed off in
HANDOFF_TO_MS_placebo_and_H2.md and results/revision/D1_placebo_characteristic_test.txt.

Output: robustness/Table_4.4a_placebo_STANDALONE.docx
"""
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "Table_4.4a_placebo_STANDALONE.docx"

doc = docx.Document()

cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = cap.add_run(
    "Table 4.4a: Placebo-Characteristic Second-Step Test, β̂_char,t ~ T_t "
    "(final table number pending the Table renumbering pass, Section E3)"
)
run.bold = True
run.font.size = Pt(10)

rows = [
    ("Characteristic", "S&P 500 HAC-t", "Full-Universe HAC-t", "Note"),
    ("Size (log mkt cap)", "−2.89", "−2.72", "significant, both panels"),
    ("Book-to-market", "+2.24", "+2.65", "significant, both panels"),
    ("Momentum (12-1)", "−2.15", "−4.00", "significant, both panels"),
    ("Market beta", "+2.61", "+1.25", "significant S&P only"),
    ("ΔS (paper's result, raw)", "+2.25", "+2.70", "reference row"),
    ("ΔS, dispersion-normalized", "+1.39", "+2.57", "S&P drops below significance; full-universe survives"),
    ("Stacked diff test, raw", "t=+2.23, p=0.026", "t=+2.26, p=0.024", "date-clustered"),
    ("Stacked diff test, normalized", "t=+1.25, p=0.213", "t=+2.21, p=0.027", "date-clustered"),
]

table = doc.add_table(rows=len(rows), cols=4)
try:
    table.style = "Table Grid"
except Exception:
    pass

for ri, row in enumerate(rows):
    for ci, val in enumerate(row):
        cell = table.cell(ri, ci)
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                if ri == 0:
                    r.bold = True

note = doc.add_paragraph()
note_run = note.add_run(
    "Source: D1 placebo-characteristic test, results/revision/D1_placebo_characteristic_test.txt. "
    "Both panels use the identical first-step regression (ret_next ~ const + ΔH_z + char_z), "
    "same controls, same months, as the paper's own ΔH/ΔS estimates (334 months S&P, "
    "111 quarters full-universe). z-scoring convention verified byte-identical to the ΔS_z "
    "construction (R20's cs_wz)."
)
note_run.italic = True
note_run.font.size = Pt(8)

doc.save(OUT)
print(f"Saved: {OUT}")

# sanity check
d2 = docx.Document(OUT)
print(f"Paragraphs: {len(d2.paragraphs)}, Tables: {len(d2.tables)}")
for r in d2.tables[0].rows:
    print([c.text for c in r.cells])
